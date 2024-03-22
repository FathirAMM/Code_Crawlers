from IPython import display
from langchain.chat_models import ChatOpenAI
from langchain.embeddings import OpenAIEmbeddings as fe
from langchain.chains import LLMChain
from langchain.prompts import PromptTemplate
import pandas as pd
from docx import Document
from docx2pdf import convert


openai_api_key = "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"

prompt_template = """
You are financial analyst and summary report writer using the providial informations.
Answer the question based only on the following context, which is from anual financial report  of srilanka 2022 and 2023.
{context}
Question: {question}
write the report with a suitable heading and use bullet points and other report wirting technics.
Answer:
"""
csv_file = "/Users/sivagar/Documents/my_projects/hackathon/report_data.csv"
docx_file = (
    "/Users/sivagar/Documents/my_projects/hackathon/economic_analysis_report.docx"
)
pdf_file = "/Users/sivagar/Documents/my_projects/hackathon/report.pdf"

question = "summarize these texts and give me the organized report"

qa_chain = LLMChain(
    llm=ChatOpenAI(model="gpt-4", openai_api_key=openai_api_key, max_tokens=3000),
    prompt=PromptTemplate.from_template(prompt_template),
)


def merge_last_two_rows(n):
    print("merging function starts")
    # Read the CSV file into a DataFrame
    df = pd.read_csv(csv_file)

    # Check if the "Response" column exists in the DataFrame
    if "Response" not in df.columns:
        raise ValueError("The 'Response' column is not present in the CSV file.")

    # Get the last two rows of the "Response" column
    last_two_rows = df["Response"].tail(n)

    # Merge the last two rows into one text
    context = last_two_rows.str.cat(sep=" ")
    print("merged finished")
    return context


def answer(context):
    print("llm function starts")
    result = qa_chain.run({"context": context, "question": question})
    return result


def create_word_document(content):
    print("creating word document start")
    # Create a new Word document
    doc = Document()

    # Split the content into sections
    sections = content.split("\n\n")

    # Add title to the document
    title = sections[0]
    doc.add_heading(title, level=1)

    # Add content to the document
    for section in sections[1:]:
        lines = section.split("\n")
        for line in lines:
            # Determine if the line is a subheading or bullet points
            if line.startswith("- "):
                # Create a bullet point
                doc.add_paragraph(line, style="ListBullet")
            else:
                # Create a regular paragraph
                doc.add_paragraph(line)

        # Add some space between sections
        doc.add_paragraph()

    # Save the Word document
    doc.save("economic_analysis_report.docx")
    print("word document created")


def convert_docx_to_pdf(docx_file, pdf_file):
    print("word to pdf starts")
    convert(docx_file, pdf_file)


def report_generation(n):
    context = merge_last_two_rows(n)
    result = answer(context)
    create_word_document(result)
    print("word document created")
    convert_docx_to_pdf(docx_file, pdf_file)
    print("pdf created")
    return True
